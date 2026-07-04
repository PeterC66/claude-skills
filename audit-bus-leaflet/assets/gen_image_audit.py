#!/usr/bin/env python
"""gen_image_audit.py — build the leaflet-image-vs-stored-data audit document.

Reads an image-audit.json file (produced by the audit-bus-leaflet skill after it has
compared one or two bus-leaflet images against the town's stored bus data) and writes a
Word .docx listing every discrepancy found, each with what the image shows, what our stored
data says, a verdict, a severity, a justification and a source. Mirrors the visual style of
make-bus-leaflet's gen_disagreements.py.

Usage:  python gen_image_audit.py <image-audit.json> [<out.docx>]
        (out.docx defaults to image-audit_<auditedOn>.docx beside the json)

image-audit.json schema:
{
  "town": "St Ives",
  "auditedOn": "2026-06-05",                         # optional; filled if absent
  "dataUsed": { "verifiedOn": "2026-06-05",          # all optional, shown in the sub-line
                "leafletVersion": "v1.0_2026-06-05",
                "files": ["verified-services.json", "routes.json", "pois.json"] },
  "images": [                                         # optional; listed in the preamble
    { "role": "internal", "title": "...", "source": "pasted|<path>" },
    { "role": "external", "title": "...", "source": "pasted|<path>" }
  ],
  "findings": [
    { "image": "external",                            # internal | external | both
      "category": "Service content",                  # Service content | POI / landmark | Spelling / consistency / design
      "item": "Route 5A (Stephensons)",
      "image_shows": "what the leaflet depicts",
      "data_says": "what our stored data says",
      "verdict": "leaflet-newer",                     # leaflet-error | leaflet-newer | data-stale | comment | ok
      "severity": "Medium",                           # High | Medium | Low | Info
      "justification": "why this matters / which side is likely right",
      "source": "verified-services.json -> notOnLeaflet[5A]" }
  ]
}
"""
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Verdict -> (row fill, human label)
VERDICT = {
    "leaflet-error": ("FCE4E4", "LEAFLET ERROR"),     # pale red  — leaflet contradicts our authoritative data
    "leaflet-newer": ("FFF2CC", "LEAFLET NEWER"),     # pale amber — leaflet ahead of our data; refresh advised
    "data-stale":    ("FFF2CC", "OUR DATA STALE"),    # pale amber — our stored data is the out-of-date side
    "comment":       ("ECECEC", "COMMENT"),           # grey      — editorial: spelling / layout / consistency
    "ok":            ("E8F4E8", "OK"),                # green     — confirmed match
}
DEFAULT_FILL = "FFFFFF"
HEADER_FILL = "2F2F2F"

SEVERITY_ORDER = {"High": 0, "Medium": 1, "Low": 2, "Info": 3}
VERDICT_ORDER = {"leaflet-error": 0, "leaflet-newer": 1, "data-stale": 2, "comment": 3, "ok": 4}
CATEGORY_ORDER = {"Service content": 0, "POI / landmark": 1, "Spelling / consistency / design": 2}


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)


def set_cell(cell, text, *, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def verdict_info(v):
    return VERDICT.get((v or "").lower(), (DEFAULT_FILL, (v or "").upper()))


def main():
    if len(sys.argv) < 2:
        sys.exit("usage: gen_image_audit.py <image-audit.json> [out.docx]")
    src = sys.argv[1]
    with open(src, "r", encoding="utf-8") as fh:
        data = json.load(fh)

    town = data.get("town", "this town")
    audited_on = data.get("auditedOn") or datetime.now().strftime("%Y-%m-%d")
    used = data.get("dataUsed", {}) or {}
    images = data.get("images", []) or []
    findings = data.get("findings", []) or []

    if len(sys.argv) > 2:
        out = sys.argv[2]
    else:
        out = os.path.join(os.path.dirname(os.path.abspath(src)),
                           f"image-audit_{audited_on}.docx")

    # Sort: by verdict (errors first), then severity, then category.
    findings = sorted(
        findings,
        key=lambda r: (
            VERDICT_ORDER.get((r.get("verdict") or "").lower(), 9),
            SEVERITY_ORDER.get(r.get("severity", "Info"), 9),
            CATEGORY_ORDER.get(r.get("category", ""), 9),
        ),
    )

    # Tallies
    counts = {}
    for r in findings:
        key = (r.get("verdict") or "").lower()
        counts[key] = counts.get(key, 0) + 1

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Pt(30)
    sec.top_margin = sec.bottom_margin = Pt(30)

    doc.add_heading(f"Bus leaflet image audit: {town}", level=0)

    # Sub-line: when audited + which stored data was used.
    bits = [f"Audited {audited_on}"]
    if used.get("verifiedOn"):
        bits.append(f"stored service data verified {used['verifiedOn']}")
    if used.get("leafletVersion"):
        bits.append(f"compared against stored leaflet {used['leafletVersion']}")
    sub = doc.add_paragraph()
    sub.add_run("   ·   ".join(bits)).italic = True

    # Count line
    order = ["leaflet-error", "leaflet-newer", "data-stale", "comment", "ok"]
    tally = "   ·   ".join(
        f"{counts[k]} {verdict_info(k)[1].lower()}" for k in order if counts.get(k)
    )
    cl = doc.add_paragraph()
    cl.add_run(f"{len(findings)} finding(s)" + (f":   {tally}" if tally else ".")).bold = True

    if used.get("files"):
        p = doc.add_paragraph()
        p.add_run("Stored data read: ").bold = True
        p.add_run(", ".join(used["files"]))

    if images:
        p = doc.add_paragraph()
        p.add_run("Images audited: ").bold = True
        p.add_run("; ".join(
            f"{im.get('role','?')} — “{im.get('title','(untitled)')}”"
            for im in images
        ))

    doc.add_paragraph(
        "Every discrepancy between the supplied leaflet image(s) and the town's stored bus "
        "data is listed below. The stored data is the reference, but the leaflet can be more "
        "current than our data — so each row carries a verdict: rows shaded red are leaflet "
        "errors (the leaflet contradicts our authoritative data); amber rows mean the leaflet "
        "looks more up-to-date than our data, or our data is the stale side (refresh advised); "
        "grey rows are editorial comments (spelling, layout, consistency); green rows confirm a "
        "match."
    )

    cols = ["Where", "Category", "Item", "Image shows", "Our stored data says",
            "Verdict", "Sev.", "Justification / source"]
    widths = [9, 16, 22, 40, 40, 16, 9, 50]  # relative hints (mm-ish)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.autofit = True
    for i, name in enumerate(cols):
        c = table.rows[0].cells[i]
        set_cell(c, name, bold=True, color="FFFFFF", size=9)
        shade(c, HEADER_FILL)

    for r in findings:
        fill, vlabel = verdict_info(r.get("verdict"))
        cells = table.add_row().cells
        set_cell(cells[0], r.get("image", ""), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[1], r.get("category", ""), size=8)
        set_cell(cells[2], r.get("item", ""), bold=True, size=9)
        set_cell(cells[3], r.get("image_shows", ""), size=8)
        set_cell(cells[4], r.get("data_says", ""), size=8)
        set_cell(cells[5], vlabel, size=8, bold=(r.get("verdict") == "leaflet-error"),
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[6], r.get("severity", ""), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        just = r.get("justification", "")
        srcref = r.get("source", "")
        tail = just + (("\n[src] " + srcref) if srcref else "")
        set_cell(cells[7], tail, size=8)
        for c in cells:
            shade(c, fill)

    # Per-category summary of the actionable rows (everything that is not 'ok').
    actionable = [r for r in findings if (r.get("verdict") or "").lower() != "ok"]
    if actionable:
        doc.add_heading("Summary by category", level=1)
        for cat in sorted({r.get("category", "Other") for r in actionable},
                          key=lambda c: CATEGORY_ORDER.get(c, 9)):
            rows_c = [r for r in actionable if r.get("category", "Other") == cat]
            doc.add_heading(f"{cat} ({len(rows_c)})", level=2)
            for r in rows_c:
                _, vlabel = verdict_info(r.get("verdict"))
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{vlabel}] {r.get('item','')} ").bold = True
                p.add_run(f"— {r.get('justification','')}")
    else:
        doc.add_paragraph("No discrepancies found — the leaflet matches our stored data.")

    doc.save(out)
    n_err = counts.get("leaflet-error", 0)
    n_refresh = counts.get("leaflet-newer", 0) + counts.get("data-stale", 0)
    print(f"wrote {out}  ({len(findings)} findings; {n_err} leaflet-error, "
          f"{n_refresh} refresh-advised)")


if __name__ == "__main__":
    main()
