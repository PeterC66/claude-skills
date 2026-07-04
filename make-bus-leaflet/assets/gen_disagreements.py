#!/usr/bin/env python
"""gen_disagreements.py — build the bustimes-vs-operator audit document (Stage S1).

Reads a disagreements.json audit file and writes a Word .docx listing EVERY route
that was checked against both bustimes.org and the operator's own website — a full
audit trail, not only the conflicts. Rows where the two sources agree are marked
"agree"; rows where they differ are shaded and the resolution taken is recorded.

Usage:  python gen_disagreements.py <disagreements.json> [<out.docx>]
        (out.docx defaults to disagreements.docx beside the json)

disagreements.json schema:
{
  "town": "March",
  "validFrom": "June 2026",          # optional
  "generatedAt": "2026-06-05T16:30",  # optional; filled if absent
  "rows": [
    { "route": "32",
      "operator": "Dews Coaches",
      "field": "operating days",      # what was compared (status / days / terminus / stops / number / fare ...)
      "bustimes": "Mon-Sat",          # what bustimes.org says
      "operator_says": "Mon-Sat",     # what the operator's site says
      "agree": true,                  # true => audit-only row; false => conflict
      "resolution": "-",              # what was done (only meaningful when agree=false)
      "sources": { "bustimes": "https://...", "operator": "https://..." }
    }
  ]
}
"""
import json
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CONFLICT_FILL = "FCE4E4"   # pale red for conflict rows
AGREE_FILL = "E8F4E8"      # pale green for agree rows
HEADER_FILL = "2F2F2F"


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)


def set_cell(cell, text, *, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def main():
    if len(sys.argv) < 2:
        sys.exit("usage: gen_disagreements.py <disagreements.json> [out.docx]")
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else (
        src.rsplit(".", 1)[0].rsplit("\\", 1)[-1].rsplit("/", 1)[-1])
    if len(sys.argv) <= 2:
        # default beside the json
        import os
        out = os.path.join(os.path.dirname(os.path.abspath(src)), "disagreements.docx")

    with open(src, "r", encoding="utf-8") as fh:
        data = json.load(fh)

    town = data.get("town", "this town")
    rows = data.get("rows", [])
    gen_at = data.get("generatedAt") or datetime.now().strftime("%Y-%m-%dT%H:%M")
    valid = data.get("validFrom", "")

    conflicts = [r for r in rows if not r.get("agree", False)]

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Pt(36)

    h = doc.add_heading(f"Bus services — bustimes.org vs operator audit: {town}", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        f"Generated {gen_at}"
        + (f"   ·   service data valid from {valid}" if valid else "")
        + f"   ·   {len(rows)} check(s), {len(conflicts)} disagreement(s)."
    ).italic = True

    doc.add_paragraph(
        "Every route checked is listed below (full audit trail). Rows shaded green agree "
        "between bustimes.org and the operator's own website. Rows shaded red disagree; the "
        "operator's site is treated as authoritative and the resolution taken is shown."
    )

    cols = ["Route", "Operator", "Field", "bustimes.org says", "Operator says", "Agree?", "Resolution / source"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.autofit = True
    for i, name in enumerate(cols):
        c = table.rows[0].cells[i]
        set_cell(c, name, bold=True, color="FFFFFF", size=9)
        shade(c, HEADER_FILL)

    for r in rows:
        agree = bool(r.get("agree", False))
        cells = table.add_row().cells
        set_cell(cells[0], r.get("route", ""), bold=True, size=9)
        set_cell(cells[1], r.get("operator", ""), size=9)
        set_cell(cells[2], r.get("field", ""), size=9)
        set_cell(cells[3], r.get("bustimes", ""), size=9)
        set_cell(cells[4], r.get("operator_says", ""), size=9)
        set_cell(cells[5], "agree" if agree else "DISAGREE", bold=not agree, size=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        srcs = r.get("sources", {}) or {}
        res = r.get("resolution", "") if not agree else ""
        tail = res
        link_bits = []
        if srcs.get("bustimes"):
            link_bits.append("bt: " + srcs["bustimes"])
        if srcs.get("operator"):
            link_bits.append("op: " + srcs["operator"])
        if link_bits:
            tail = (res + "\n" if res else "") + "\n".join(link_bits)
        set_cell(cells[6], tail, size=8)
        fill = AGREE_FILL if agree else CONFLICT_FILL
        for c in cells:
            shade(c, fill)

    if conflicts:
        doc.add_heading("Disagreements summary", level=1)
        for r in conflicts:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{r.get('route','')} ({r.get('operator','')}) — {r.get('field','')}: ").bold = True
            p.add_run(f"bustimes.org: {r.get('bustimes','')}; operator: {r.get('operator_says','')}. ")
            p.add_run(f"Resolution: {r.get('resolution','')}.").italic = True
    else:
        doc.add_paragraph("No disagreements found — bustimes.org and every operator site agreed.")

    doc.save(out)
    print(f"wrote {out}  ({len(rows)} rows, {len(conflicts)} disagreement(s))")


if __name__ == "__main__":
    main()
