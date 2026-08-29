#!/usr/bin/env python
"""gen_verification.py — render the S6 independent-verification report to .docx.

Reads a verification.json (produced by verify_report.js — the antagonistic
verification pass: a blind red-team re-derivation diffed against our pipeline,
plus structural/geographic sanity checks) and writes a Word .docx reliability
report. Every finding is listed and classified HARD (blocks the build) or SOFT
(logged only). Style mirrors gen_disagreements.py.

Usage:  python gen_verification.py <verification.json> [<out.docx>]
        (out.docx defaults to verification.docx beside the json)

verification.json schema (see verify_report.js):
{
  "town": "St Ives",
  "generatedAt": "2026-06-07T15:01",
  "redteamPresent": true,
  "redteamSources": ["https://..."],
  "inputs": {"verifiedOn": "...", "routesVersion": "4.0", "displayedRoutes": [...]},
  "summary": {"checks": N, "hard": N, "soft": N, "pass": true},
  "findings": [
    {"id": "F001", "severity": "hard|soft", "category": "...", "route": "301",
     "message": "...", "evidence": {...}, "source": "sanity|redteam"}
  ]
}
"""
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HARD_FILL = "FCE4E4"     # pale red for hard findings
SOFT_FILL = "FFF4D6"     # pale amber for soft findings
HEADER_FILL = "2F2F2F"
PASS_FILL = "E8F4E8"     # pale green banner
BLOCK_FILL = "F4C7C7"    # stronger red banner


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)


def set_cell(cell, text, *, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_col_widths(table, usable_width, weights):
    """Fix each column to a proportional share of the page's usable width —
    see gen_disagreements.py's copy of this helper for why `table.autofit`
    alone isn't enough (it only autofits when Word itself opens the file;
    the PDF is rendered by headless LibreOffice, which just splits the grid
    evenly, cramping Finding/Evidence against ID/Severity/Route)."""
    widths = [Emu(int(usable_width * w / sum(weights))) for w in weights]
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
    # Per-cell tcW alone isn't enough: python-docx leaves tblGrid at the
    # equal-width columns it created the table with, and that's what decides
    # layout wherever no cell in a column happens to override it (e.g. a
    # short table with fewer rows than columns wide). Rewrite it to match.
    grid = table._tbl.find(qn("w:tblGrid"))
    for gridcol, w in zip(grid.findall(qn("w:gridCol")), widths):
        gridcol.set(qn("w:w"), str(w.twips))


def evidence_str(ev):
    if not ev:
        return ""
    bits = []
    for k, v in ev.items():
        if isinstance(v, (dict, list)):
            v = json.dumps(v, ensure_ascii=False, separators=(",", ":"))
        bits.append(f"{k}: {v}")
    return "\n".join(bits)


def main():
    if len(sys.argv) < 2:
        sys.exit("usage: gen_verification.py <verification.json> [out.docx]")
    src = sys.argv[1]
    if len(sys.argv) > 2:
        out = sys.argv[2]
    else:
        out = os.path.join(os.path.dirname(os.path.abspath(src)), "verification.docx")

    with open(src, "r", encoding="utf-8") as fh:
        data = json.load(fh)

    town = data.get("town", "this town")
    findings = data.get("findings", [])
    summary = data.get("summary", {})
    inputs = data.get("inputs", {})
    gen_at = data.get("generatedAt") or datetime.now().strftime("%Y-%m-%dT%H:%M")
    n_hard = summary.get("hard", sum(1 for f in findings if f.get("severity") == "hard"))
    n_soft = summary.get("soft", sum(1 for f in findings if f.get("severity") == "soft"))
    passed = summary.get("pass", n_hard == 0)
    rt = data.get("redteamPresent", False)

    hard = [f for f in findings if f.get("severity") == "hard"]
    soft = [f for f in findings if f.get("severity") == "soft"]

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Pt(36)

    doc.add_heading(f"Bus services — independent verification report: {town}", level=0)

    sub = doc.add_paragraph()
    sub.add_run(
        f"Generated {gen_at}"
        + (f"   ·   routes v{inputs.get('routesVersion')}" if inputs.get("routesVersion") else "")
        + (f"   ·   services verified {inputs.get('verifiedOn')}" if inputs.get("verifiedOn") else "")
        + f"   ·   {n_hard} hard, {n_soft} soft finding(s)."
    ).italic = True

    # PASS / BLOCKED banner
    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    bc = banner.rows[0].cells[0]
    set_cell(bc,
             ("RESULT: PASS — no blocking findings; the stored data is safe to build/rely on."
              if passed else
              "RESULT: BLOCKED — hard findings must be resolved before this build can be trusted."),
             bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=("1E5E2E" if passed else "8A1C1C"))
    shade(bc, PASS_FILL if passed else BLOCK_FILL)

    doc.add_paragraph(
        "This is the antagonistic / independent verification pass (Stage S6). "
        + ("An independent blind red-team agent re-derived the town's services from scratch "
           "(operator, termini, operating days, and whether each route serves the town) using "
           "bustimes.org plus a second source, with no sight of our stored data; its findings "
           "were then diffed against our pipeline. " if rt else
           "NOTE: no red-team file was present, so only the structural / geographic sanity "
           "checks ran. ")
        + "In addition, structural and geographic sanity checks were run against the stored "
        "geometry. Findings are classified HARD (would make the leaflet wrong or undrawable — "
        "these block the build) or SOFT (naming, day variants, off-by-one, inclusion candidates "
        "— logged for review only)."
    )
    if rt and data.get("redteamSources"):
        p = doc.add_paragraph()
        p.add_run("Red-team sources: ").bold = True
        p.add_run("; ".join(data["redteamSources"])).italic = True
    if inputs.get("displayedRoutes"):
        p = doc.add_paragraph()
        p.add_run("Routes drawn on the leaflet: ").bold = True
        p.add_run(", ".join(inputs["displayedRoutes"]))

    # findings table (hard first, then soft)
    cols = ["ID", "Severity", "Category", "Route", "Finding", "Evidence / source"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.autofit = True
    for i, name in enumerate(cols):
        c = table.rows[0].cells[i]
        set_cell(c, name, bold=True, color="FFFFFF", size=9)
        shade(c, HEADER_FILL)

    if not findings:
        cells = table.add_row().cells
        set_cell(cells[0], "—", size=9)
        set_cell(cells[4], "No findings — every check was clean.", size=9)
        for c in cells:
            shade(c, PASS_FILL)

    for f in hard + soft:
        is_hard = f.get("severity") == "hard"
        cells = table.add_row().cells
        set_cell(cells[0], f.get("id", ""), bold=True, size=9)
        set_cell(cells[1], "HARD" if is_hard else "soft", bold=is_hard, size=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=("8A1C1C" if is_hard else None))
        set_cell(cells[2], f.get("category", ""), size=9)
        set_cell(cells[3], f.get("route") or "", bold=True, size=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(cells[4], f.get("message", ""), size=9)
        tail = evidence_str(f.get("evidence"))
        srctag = f.get("source", "")
        tail = (f"[{srctag}]\n" if srctag else "") + tail
        set_cell(cells[5], tail, size=7)
        fill = HARD_FILL if is_hard else SOFT_FILL
        for c in cells:
            shade(c, fill)

    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    # ID / Severity / Route are short codes; Finding is the prose message
    # and Evidence/source carries raw JSON, so they take the largest shares.
    set_col_widths(table, usable, [0.5, 0.8, 1.0, 0.6, 2.6, 2.0])

    # summary sections
    if hard:
        doc.add_heading("Hard findings — must resolve before relying on this build", level=1)
        for f in hard:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{f.get('id','')}] {f.get('category','')}"
                      + (f" {f.get('route')}" if f.get("route") else "") + ": ").bold = True
            p.add_run(f.get("message", ""))
    if soft:
        doc.add_heading("Soft findings — review (not blocking)", level=1)
        for f in soft:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{f.get('id','')}] {f.get('category','')}"
                      + (f" {f.get('route')}" if f.get("route") else "") + ": ").bold = True
            p.add_run(f.get("message", ""))
    if not findings:
        doc.add_paragraph("No findings — the independent pass and the sanity checks all agreed "
                          "with the stored data.")

    # Stamp real created/modified dates (python-docx's blank template otherwise
    # leaves its 2013-12-23 date, which Explorer shows and reads as wrong).
    _now = datetime.now()
    doc.core_properties.created = _now
    doc.core_properties.modified = _now
    doc.save(out)
    print(f"wrote {out}  ({n_hard} hard, {n_soft} soft, {'PASS' if passed else 'BLOCKED'})")


if __name__ == "__main__":
    main()
